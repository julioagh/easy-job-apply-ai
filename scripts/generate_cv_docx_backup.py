#!/usr/bin/env python3
"""
Generador de CV en formato DOCX optimizado para ATS
Sistema: Easy Job Apply AI v2.0
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=5, is_last_in_section=False):
    """
    Configura formato estándar del párrafo

    Args:
        paragraph: Objeto paragraph de python-docx
        alignment: Tipo de alineación (JUSTIFY, CENTER, LEFT)
        space_after: Espacio después en puntos (default 5pt)
        is_last_in_section: Si es última línea de sección (usa 10pt)
    """
    # Interlineado simple (1.0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Espacio después: 5pt normal, 10pt si es última de sección
    paragraph.paragraph_format.space_after = Pt(10 if is_last_in_section else space_after)

    # Alineación
    paragraph.alignment = alignment

def create_cv_docx():
    """Genera CV en formato DOCX con formato profesional ATS-friendly"""

    # Crear documento
    doc = Document()

    # Configurar márgenes exactos (1.22 cm)
    section = doc.sections[0]
    section.top_margin = Cm(1.22)
    section.bottom_margin = Cm(1.22)
    section.left_margin = Cm(1.22)
    section.right_margin = Cm(1.22)

    # ========== HEADER: NOMBRE Y CONTACTO ==========

    # Nombre (18pt Bold) - CENTRADO
    nombre = doc.add_paragraph()
    run = nombre.add_run('JULIO ALBERTO GONZALES HEREDIA')
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    set_paragraph_format(nombre, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # Contacto (9pt Normal) - CENTRADO
    contacto = doc.add_paragraph()
    run = contacto.add_run('Lima, Perú | +51 992755873 | jgonzales.sbs@gmail.com | LinkedIn: linkedin.com/in/julioagh')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    set_paragraph_format(contacto, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5, is_last_in_section=True)

    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # ========== RESUMEN PROFESIONAL ==========

    # Header sección (12pt Bold UPPERCASE) - JUSTIFICADO
    header = doc.add_paragraph()
    run = header.add_run('RESUMEN PROFESIONAL')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    set_paragraph_format(header, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=5)

    # Contenido summary (9.5pt Normal) - JUSTIFICADO, última de sección
    summary = doc.add_paragraph()
    run = summary.add_run('Agile Coach Expert con 7 años liderando transformación organizacional y agilidad escalada (SAFe, Scrum@Scale) en multinacionales de banca y telecom. Especializado en implementación de OKRs, métricas ágiles (Velocity, Lead Time, Cycle Time, Throughput), Change Management y facilitación masiva. Certificado AI+ Executive (2025) con experiencia aplicando IA generativa en contextos organizacionales. Experto en coaching de equipos, mentoring de líderes, gestión de backlog y transformación digital. Track record comprobado gestionando transformaciones a escala (100+ equipos/proyectos), actuando como Scrum Master senior y liderando formación corporativa con enfoque data-driven.')
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    set_paragraph_format(summary, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=5, is_last_in_section=True)

    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # ========== EXPERIENCIA PROFESIONAL ==========

    # Header sección
    header = doc.add_paragraph()
    run = header.add_run('EXPERIENCIA PROFESIONAL')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    # BBVA PERÚ
    empresa = doc.add_paragraph()
    run = empresa.add_run('BBVA PERÚ | Lima, Perú')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = False  # NO bold

    cargo = doc.add_paragraph()
    run = cargo.add_run('Agile Coach Expert | Enero 2022 - Presente')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = False

    # Bullets BBVA (9pt Normal)
    bullets_bbva = [
        'Lideré transformación organizacional a escala en 4 VPs estratégicas gestionando 100+ equipos/proyectos, implementando agilidad escalada (SAFe, Scrum@Scale) con foco en OKRs y métricas de flujo, logrando alineación estratégica corporativa y mejora en predictibilidad de entrega',
        'Diseñé e implementé modelo de governance con dominios, roles, responsabilidades y KPIs de efectividad organizacional, estableciendo métricas ágiles (Velocity, Cycle Time, Lead Time, Throughput) conectando capacidad de equipos con objetivos empresariales mediante dashboards ejecutivos',
        'Lideré iniciativa global de Predictibilidad y Efectividad de Entregables como referente senior, diseñando marcos de métricas aplicables a LATAM y coordinando con equipos regionales en gestión de backlog y flow analytics para decisiones data-driven',
        'Representé a Perú (6 meses) en iniciativa regional de Change Management BBVA hispanohablante, co-diseñando modelo de gestión del cambio cultural alineado con estándares globales y liderando prácticas de transformación organizacional',
        'Dirigí diseño e implementación de formación en agilidad para TODO el banco (3 años) como facilitador principal y mentor, entrenando 450+ personas virtualmente y 80+ presencialmente en Scrum, coaching y metodologías ágiles',
        'Actué como mentor de líderes en planificaciones trimestrales y alineamiento organizacional, liderando ruta formativa de Agile Coaches hispanohablantes en product strategy, métricas ágiles y facilitación avanzada',
        'Gestioné delivery de hasta 7 equipos/proyectos simultáneamente como Scrum Master, asegurando alineamiento de expectativas de stakeholders, gestión de backlog, dependencias, riesgos y calidad en la entrega'
    ]

    for bullet_text in bullets_bbva:
        bullet = doc.add_paragraph()
        run = bullet.add_run('• ' + bullet_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Espacio entre roles
    doc.add_paragraph()

    # ENTEL PERÚ
    empresa = doc.add_paragraph()
    run = empresa.add_run('ENTEL PERÚ | Lima, Perú')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = False

    cargo = doc.add_paragraph()
    run = cargo.add_run('Agile Coach | Julio 2018 - Diciembre 2022')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = False

    # Bullets Entel
    bullets_entel = [
        'Lideré 2 fases de transformación ágil corporativa multi-fase con cambios estructurales profundos, implementando agilidad escalada (modelo Spotify adaptado) y coordinando con matriz Chile para estandarización grupal en prácticas Scrum y Kanban',
        'Diseñé nuevas estructuras organizacionales con área Talento: perfiles, roles, responsabilidades y cargas de trabajo alineadas a transformación estratégica, estableciendo modelo presupuestario/financiero con mayor autonomía',
        'Implementé suite completa de OKRs y métricas de desempeño organizacional conectando objetivos de equipos con estrategia corporativa, logrando 20% de ahorro en costos operativos y 5% de mejora en time-to-market',
        'Definí indicadores clave para evaluar adopción de prácticas ágiles y salud organizacional, implementando métricas ágiles (Velocity, Lead Time, Throughput) y flow analytics para toma de decisiones data-driven',
        'Lideré iniciativa de Change Management cultural mediante talleres, comunidades de práctica, coaching y mentoring, gestionando resistencia organizacional en contexto de transformación profunda',
        'Proporcioné asesoría estratégica a líderes senior en cambio organizacional, gestión de backlog y transformación de cultura, actuando como Scrum Master de referencia y facilitador principal',
        'Gestioné delivery de hasta 9 equipos/proyectos simultáneamente, asegurando alineamiento de stakeholders, gestión de dependencias y calidad de entregables'
    ]

    for bullet_text in bullets_entel:
        bullet = doc.add_paragraph()
        run = bullet.add_run('• ' + bullet_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Espacio entre roles
    doc.add_paragraph()

    # TRANSFORMACIÓN DIGITAL 2007-2017
    empresa = doc.add_paragraph()
    run = empresa.add_run('TRANSFORMACIÓN DIGITAL & MODELOS OPERATIVOS | Lima, Perú')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = False

    cargo = doc.add_paragraph()
    run = cargo.add_run('Agile Coach, DevOps Specialist, Scrum Master | 2007 - 2017')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = False

    subempresa = doc.add_paragraph()
    run = subempresa.add_run('SOAINT Consulting | HP Perú | Novatronic/Agile Works')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.italic = True

    bullet = doc.add_paragraph()
    run = bullet.add_run('• Lideré iniciativas de transformación digital, implementación DevOps y CI/CD, coaching de equipos multidisciplinarios en metodologías ágiles (Scrum, Kanban) y gestión de backlog en proyectos de desarrollo de software')
    run.font.name = 'Arial'
    run.font.size = Pt(9)

    # Línea separadora
    doc.add_paragraph('_' * 90)

    # ========== FORMACIÓN Y CERTIFICACIONES ==========

    header = doc.add_paragraph()
    run = header.add_run('FORMACIÓN Y CERTIFICACIONES')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    # Formación Ejecutiva (2025)
    subheader = doc.add_paragraph()
    run = subheader.add_run('Formación Ejecutiva (2025)')
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.bold = True

    certs_ejecutiva = [
        'AI+ Executive (AICerts) - Estrategias GenAI, IA generativa, analítica predictiva y ROI en contextos organizacionales; incluye People Analytics & Machine Learning aplicado a casos de uso organizacionales',
        'Organization Design Certificate (Org-ology - Accredited CODP pathway) - Metodologías de diagnóstico, diseño y rediseño organizacional',
        'Strategy & Execution Specialist Program (UMAAN) - Framework sistémico para definición de estrategia y ejecución'
    ]

    for cert in certs_ejecutiva:
        p = doc.add_paragraph()
        run = p.add_run('• ' + cert)
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Certificaciones Agile & Coaching
    subheader = doc.add_paragraph()
    run = subheader.add_run('Certificaciones Agile & Coaching')
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.bold = True

    certs_agile = [
        'SAFe Program Consultant (SPC) - Scaled Agile Framework para agilidad escalada',
        'Scrum@Scale Practitioner - Framework de escalamiento Scrum',
        'Coaching Agile Teams (CAT 1 & 2) - Coaching avanzado de equipos ágiles',
        'Certified ScrumMaster (CSM) | Certified Scrum Product Owner (CSPO) | Advanced Certified Scrum Developer (A-CSD)',
        'Lean Change Agent - Change Management en transformaciones organizacionales',
        'Kanban Coaching Professional (KCP)',
        'Certified Lean Inception Facilitator'
    ]

    for cert in certs_agile:
        p = doc.add_paragraph()
        run = p.add_run('• ' + cert)
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Certificaciones Liderazgo & Facilitación
    subheader = doc.add_paragraph()
    run = subheader.add_run('Certificaciones Liderazgo & Facilitación')
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.bold = True

    p = doc.add_paragraph()
    run = p.add_run('• Leadership Circle Profile (2018) | Conscious Business Coach | Real L.I.F.E. Facilitator | Lego Serious Play Facilitator | MBTI (Step I & II)')
    run.font.name = 'Arial'
    run.font.size = Pt(9)

    # Formación Académica
    subheader = doc.add_paragraph()
    run = subheader.add_run('Formación Académica')
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.bold = True

    p = doc.add_paragraph()
    run = p.add_run('• Bachiller en Ciencias con mención en Ingeniería Informática - PUCP, 2007')
    run.font.name = 'Arial'
    run.font.size = Pt(9)

    # Línea separadora
    doc.add_paragraph('_' * 90)

    # ========== COMPETENCIAS CLAVE ==========

    header = doc.add_paragraph()
    run = header.add_run('COMPETENCIAS CLAVE')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    competencias = [
        ('Agilidad Escalada & Frameworks:', 'SAFe, Scrum@Scale, Spotify Model, Scrum, Kanban, Lean, Design Thinking'),
        ('Gestión & Métricas:', 'OKRs, KPIs organizacionales, métricas ágiles (Velocity, Cycle Time, Lead Time, Throughput), flow analytics, gestión de backlog, data-driven decision making'),
        ('Transformación Organizacional:', 'Change Management, gestión del cambio cultural, transformación digital, diseño organizacional, governance'),
        ('Liderazgo & Desarrollo:', 'Coaching (Agile Coach, team coaching), mentoring, facilitación de workshops masivos (hasta 450 personas), formación corporativa, Scrum Master senior'),
        ('IA Generativa & Tecnología:', 'IA generativa, Claude, API Integration, Agentes de IA, Machine Learning, People Analytics, casos de uso organizacionales (L&D, Product Strategy)'),
        ('Gestión de Proyectos:', 'Scrum Master (7+ años), Delivery Management, gestión de dependencias y riesgos, alineamiento de stakeholders'),
        ('Herramientas:', 'Jira, Azure DevOps, Confluence, Miro, herramientas de gestión de proyectos ágiles')
    ]

    for categoria, items in competencias:
        p = doc.add_paragraph()
        # Categoría en bold
        run = p.add_run(categoria)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.bold = True
        # Items en normal
        run = p.add_run(' ' + items)
        run.font.name = 'Arial'
        run.font.size = Pt(9)

    # Línea separadora
    doc.add_paragraph('_' * 90)

    # ========== IDIOMAS ==========

    header = doc.add_paragraph()
    run = header.add_run('IDIOMAS')
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    idiomas = doc.add_paragraph()
    run = idiomas.add_run('Español - Nativo | Inglés - Avanzado (C1)')
    run.font.name = 'Arial'
    run.font.size = Pt(9)

    # Guardar documento
    output_path = '/sessions/optimistic-determined-hopper/mnt/easy-job-apply-ai/outputs/CV_Gonzales_AttachGroup_AgileCoach.docx'
    doc.save(output_path)

    return output_path

if __name__ == '__main__':
    output = create_cv_docx()
    print(f"✅ CV generado exitosamente: {output}")
