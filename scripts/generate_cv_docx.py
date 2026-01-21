#!/usr/bin/env python3
"""
Generador de CV en formato DOCX optimizado para ATS
Sistema: Easy Job Apply AI v2.0
Versión: 2.0 - Con rutas configurables y mejoras estructurales
"""

import sys
from pathlib import Path

# Agregar directorio scripts al path para importar config
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from config import (
    get_output_path,
    CV_MARGINS_CM,
    CV_FONT_SIZES,
    DEFAULT_FONT,
    CANDIDATE_INFO
)

def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=5, is_last_in_section=False):
    """
    Configura formato estándar del párrafo
    
    Args:
        paragraph: Objeto paragraph de python-docx
        alignment: Tipo de alineación (JUSTIFY, CENTER, LEFT)
        space_after: Espacio después en puntos (default 5pt)
        is_last_in_section: Si es última línea de sección (usa 10pt)
    """
    # Interlineado simple (1.0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Espacio después: 5pt normal, 10pt si es última de sección
    paragraph.paragraph_format.space_after = Pt(10 if is_last_in_section else space_after)
    
    # Alineación
    paragraph.alignment = alignment

def add_text_run(paragraph, text, font_size, bold=False, italic=False, font_name=None):
    """Helper para agregar texto con formato"""
    run = paragraph.add_run(text)
    run.font.name = font_name or DEFAULT_FONT
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    return run

def create_cv_docx(output_filename='CV_Gonzales_AttachGroup_AgileCoach.docx',
                   candidate_info=None):
    """
    Genera CV en formato DOCX con formato profesional ATS-friendly

    Args:
        output_filename: Nombre del archivo de salida (default: CV_Gonzales_AttachGroup_AgileCoach.docx)
        candidate_info: Dict con información del candidato (opcional, usa config por defecto)

    Returns:
        Path: Ruta completa del archivo generado
    """
    # Usar información del candidato de config si no se proporciona
    if candidate_info is None:
        candidate_info = CANDIDATE_INFO

    # Crear documento
    doc = Document()

    # Configurar márgenes desde config
    section = doc.sections[0]
    section.top_margin = Cm(CV_MARGINS_CM)
    section.bottom_margin = Cm(CV_MARGINS_CM)
    section.left_margin = Cm(CV_MARGINS_CM)
    section.right_margin = Cm(CV_MARGINS_CM)
    
    # ========== HEADER: NOMBRE Y CONTACTO (CENTRADO) ==========

    # Nombre (18pt Bold)
    nombre = doc.add_paragraph()
    add_text_run(nombre, candidate_info['nombre'], CV_FONT_SIZES['nombre'], bold=True)
    set_paragraph_format(nombre, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # Contacto (9pt Normal)
    contacto = doc.add_paragraph()
    contacto_text = f"{candidate_info['ubicacion']} | {candidate_info['telefono']} | {candidate_info['email']} | LinkedIn: {candidate_info['linkedin']}"
    add_text_run(contacto, contacto_text, CV_FONT_SIZES['contacto'])
    set_paragraph_format(contacto, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5, is_last_in_section=True)
    
    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # ========== RESUMEN PROFESIONAL (JUSTIFICADO) ==========
    
    # Header sección (12pt Bold UPPERCASE)
    header = doc.add_paragraph()
    add_text_run(header, 'RESUMEN PROFESIONAL', CV_FONT_SIZES['header'], bold=True)
    set_paragraph_format(header, space_after=5)

    # Contenido summary (9.5pt Normal)
    summary = doc.add_paragraph()
    add_text_run(summary, 'Agile Coach Expert con 7 años liderando transformación organizacional y agilidad escalada (SAFe, Scrum@Scale) en multinacionales de banca y telecom. Especializado en implementación de OKRs, métricas ágiles (Velocity, Lead Time, Cycle Time, Throughput), Change Management y facilitación masiva. Certificado AI+ Executive (2025) con experiencia aplicando IA generativa en contextos organizacionales. Experto en coaching de equipos, mentoring de líderes, gestión de backlog y transformación digital. Track record comprobado gestionando transformaciones a escala (100+ equipos/proyectos), actuando como Scrum Master senior y liderando formación corporativa con enfoque data-driven.', CV_FONT_SIZES['summary'])
    set_paragraph_format(summary, space_after=5, is_last_in_section=True)
    
    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # ========== EXPERIENCIA PROFESIONAL (JUSTIFICADO) ==========
    
    # Header sección
    header = doc.add_paragraph()
    add_text_run(header, 'EXPERIENCIA PROFESIONAL', CV_FONT_SIZES['header'], bold=True)
    set_paragraph_format(header, space_after=5)
    
    # BBVA PERÚ
    empresa = doc.add_paragraph()
    add_text_run(empresa, 'BBVA PERÚ | Lima, Perú', CV_FONT_SIZES['body'])
    set_paragraph_format(empresa, space_after=5)
    
    cargo = doc.add_paragraph()
    add_text_run(cargo, 'Agile Coach Expert | Enero 2022 - Presente', CV_FONT_SIZES['body'])
    set_paragraph_format(cargo, space_after=5)
    
    # Bullets BBVA
    bullets_bbva = [
        'Lideré transformación organizacional a escala en 4 VPs estratégicas gestionando 100+ equipos/proyectos, implementando agilidad escalada (SAFe, Scrum@Scale) con foco en OKRs y métricas de flujo, logrando alineación estratégica corporativa y mejora en predictibilidad de entrega',
        'Diseñé e implementé modelo de governance con dominios, roles, responsabilidades y KPIs de efectividad organizacional, estableciendo métricas ágiles (Velocity, Cycle Time, Lead Time, Throughput) conectando capacidad de equipos con objetivos empresariales mediante dashboards ejecutivos',
        'Lideré iniciativa global de Predictibilidad y Efectividad de Entregables como referente senior, diseñando marcos de métricas aplicables a LATAM y coordinando con equipos regionales en gestión de backlog y flow analytics para decisiones data-driven',
        'Representé a Perú (6 meses) en iniciativa regional de Change Management BBVA hispanohablante, co-diseñando modelo de gestión del cambio cultural alineado con estándares globales y liderando prácticas de transformación organizacional',
        'Dirigí diseño e implementación de formación en agilidad para TODO el banco (3 años) como facilitador principal y mentor, entrenando 450+ personas virtualmente y 80+ presencialmente en Scrum, coaching y metodologías ágiles',
        'Actué como mentor de líderes en planificaciones trimestrales y alineamiento organizacional, liderando ruta formativa de Agile Coaches hispanohablantes en product strategy, métricas ágiles y facilitación avanzada',
        'Gestioné delivery de hasta 7 equipos/proyectos simultáneamente como Scrum Master, asegurando alineamiento de expectativas de stakeholders, gestión de backlog, dependencias, riesgos y calidad en la entrega'
    ]
    
    for i, bullet_text in enumerate(bullets_bbva):
        is_last = (i == len(bullets_bbva) - 1)
        bullet = doc.add_paragraph()
        add_text_run(bullet, '• ' + bullet_text, CV_FONT_SIZES['body'])
        set_paragraph_format(bullet, space_after=5, is_last_in_section=is_last)
    
    # ENTEL PERÚ
    empresa = doc.add_paragraph()
    add_text_run(empresa, 'ENTEL PERÚ | Lima, Perú', CV_FONT_SIZES['body'])
    set_paragraph_format(empresa, space_after=5)
    
    cargo = doc.add_paragraph()
    add_text_run(cargo, 'Agile Coach | Julio 2018 - Diciembre 2022', CV_FONT_SIZES['body'])
    set_paragraph_format(cargo, space_after=5)
    
    # Bullets Entel
    bullets_entel = [
        'Lideré 2 fases de transformación ágil corporativa multi-fase con cambios estructurales profundos, implementando agilidad escalada (modelo Spotify adaptado) y coordinando con matriz Chile para estandarización grupal en prácticas Scrum y Kanban',
        'Diseñé nuevas estructuras organizacionales con área Talento: perfiles, roles, responsabilidades y cargas de trabajo alineadas a transformación estratégica, estableciendo modelo presupuestario/financiero con mayor autonomía',
        'Implementé suite completa de OKRs y métricas de desempeño organizacional conectando objetivos de equipos con estrategia corporativa, logrando 20% de ahorro en costos operativos y 5% de mejora en time-to-market',
        'Definí indicadores clave para evaluar adopción de prácticas ágiles y salud organizacional, implementando métricas ágiles (Velocity, Lead Time, Throughput) y flow analytics para toma de decisiones data-driven',
        'Lideré iniciativa de Change Management cultural mediante talleres, comunidades de práctica, coaching y mentoring, gestionando resistencia organizacional en contexto de transformación profunda',
        'Proporcioné asesoría estratégica a líderes senior en cambio organizacional, gestión de backlog y transformación de cultura, actuando como Scrum Master de referencia y facilitador principal',
        'Gestioné delivery de hasta 9 equipos/proyectos simultáneamente, asegurando alineamiento de stakeholders, gestión de dependencias y calidad de entregables'
    ]
    
    for i, bullet_text in enumerate(bullets_entel):
        is_last = (i == len(bullets_entel) - 1)
        bullet = doc.add_paragraph()
        add_text_run(bullet, '• ' + bullet_text, CV_FONT_SIZES['body'])
        set_paragraph_format(bullet, space_after=5, is_last_in_section=is_last)
    
    # TRANSFORMACIÓN DIGITAL 2007-2017
    empresa = doc.add_paragraph()
    add_text_run(empresa, 'TRANSFORMACIÓN DIGITAL & MODELOS OPERATIVOS | Lima, Perú', CV_FONT_SIZES['body'])
    set_paragraph_format(empresa, space_after=5)
    
    cargo = doc.add_paragraph()
    add_text_run(cargo, 'Agile Coach, DevOps Specialist, Scrum Master | 2007 - 2017', CV_FONT_SIZES['body'])
    set_paragraph_format(cargo, space_after=5)
    
    subempresa = doc.add_paragraph()
    add_text_run(subempresa, 'SOAINT Consulting | HP Perú | Novatronic/Agile Works', 9, italic=True)
    set_paragraph_format(subempresa, space_after=5)
    
    bullet = doc.add_paragraph()
    add_text_run(bullet, '• Lideré iniciativas de transformación digital, implementación DevOps y CI/CD, coaching de equipos multidisciplinarios en metodologías ágiles (Scrum, Kanban) y gestión de backlog en proyectos de desarrollo de software', CV_FONT_SIZES['body'])
    set_paragraph_format(bullet, space_after=5, is_last_in_section=True)
    
    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # ========== FORMACIÓN Y CERTIFICACIONES (JUSTIFICADO) ==========
    
    header = doc.add_paragraph()
    add_text_run(header, 'FORMACIÓN Y CERTIFICACIONES', CV_FONT_SIZES['header'], bold=True)
    set_paragraph_format(header, space_after=5)
    
    # Formación Ejecutiva (2025)
    subheader = doc.add_paragraph()
    add_text_run(subheader, 'Formación Ejecutiva (2025)', CV_FONT_SIZES['subheader'], bold=True)
    set_paragraph_format(subheader, space_after=5)
    
    certs_ejecutiva = [
        'AI+ Executive (AICerts) - Estrategias GenAI, IA generativa, analítica predictiva y ROI en contextos organizacionales; incluye People Analytics & Machine Learning aplicado a casos de uso organizacionales',
        'Organization Design Certificate (Org-ology - Accredited CODP pathway) - Metodologías de diagnóstico, diseño y rediseño organizacional',
        'Strategy & Execution Specialist Program (UMAAN) - Framework sistémico para definición de estrategia y ejecución'
    ]
    
    for i, cert in enumerate(certs_ejecutiva):
        is_last = (i == len(certs_ejecutiva) - 1)
        p = doc.add_paragraph()
        add_text_run(p, '• ' + cert, CV_FONT_SIZES['body'])
        set_paragraph_format(p, space_after=5, is_last_in_section=is_last)
    
    # Certificaciones Agile & Coaching
    subheader = doc.add_paragraph()
    add_text_run(subheader, 'Certificaciones Agile & Coaching', CV_FONT_SIZES['subheader'], bold=True)
    set_paragraph_format(subheader, space_after=5)
    
    certs_agile = [
        'SAFe Program Consultant (SPC) - Scaled Agile Framework para agilidad escalada',
        'Scrum@Scale Practitioner - Framework de escalamiento Scrum',
        'Coaching Agile Teams (CAT 1 & 2) - Coaching avanzado de equipos ágiles',
        'Certified ScrumMaster (CSM) | Certified Scrum Product Owner (CSPO) | Advanced Certified Scrum Developer (A-CSD)',
        'Lean Change Agent - Change Management en transformaciones organizacionales',
        'Kanban Coaching Professional (KCP)',
        'Certified Lean Inception Facilitator'
    ]
    
    for i, cert in enumerate(certs_agile):
        is_last = (i == len(certs_agile) - 1)
        p = doc.add_paragraph()
        add_text_run(p, '• ' + cert, CV_FONT_SIZES['body'])
        set_paragraph_format(p, space_after=5, is_last_in_section=is_last)
    
    # Certificaciones Liderazgo & Facilitación
    subheader = doc.add_paragraph()
    add_text_run(subheader, 'Certificaciones Liderazgo & Facilitación', CV_FONT_SIZES['subheader'], bold=True)
    set_paragraph_format(subheader, space_after=5)
    
    p = doc.add_paragraph()
    add_text_run(p, '• Leadership Circle Profile (2018) | Conscious Business Coach | Real L.I.F.E. Facilitator | Lego Serious Play Facilitator | MBTI (Step I & II)', CV_FONT_SIZES['body'])
    set_paragraph_format(p, space_after=5, is_last_in_section=True)
    
    # Formación Académica
    subheader = doc.add_paragraph()
    add_text_run(subheader, 'Formación Académica', CV_FONT_SIZES['subheader'], bold=True)
    set_paragraph_format(subheader, space_after=5)
    
    p = doc.add_paragraph()
    add_text_run(p, '• Bachiller en Ciencias con mención en Ingeniería Informática - PUCP, 2007', CV_FONT_SIZES['body'])
    set_paragraph_format(p, space_after=5, is_last_in_section=True)
    
    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # ========== COMPETENCIAS CLAVE (JUSTIFICADO) ==========
    
    header = doc.add_paragraph()
    add_text_run(header, 'COMPETENCIAS CLAVE', CV_FONT_SIZES['header'], bold=True)
    set_paragraph_format(header, space_after=5)
    
    competencias = [
        ('Agilidad Escalada & Frameworks:', 'SAFe, Scrum@Scale, Spotify Model, Scrum, Kanban, Lean, Design Thinking'),
        ('Gestión & Métricas:', 'OKRs, KPIs organizacionales, métricas ágiles (Velocity, Cycle Time, Lead Time, Throughput), flow analytics, gestión de backlog, data-driven decision making'),
        ('Transformación Organizacional:', 'Change Management, gestión del cambio cultural, transformación digital, diseño organizacional, governance'),
        ('Liderazgo & Desarrollo:', 'Coaching (Agile Coach, team coaching), mentoring, facilitación de workshops masivos (hasta 450 personas), formación corporativa, Scrum Master senior'),
        ('IA Generativa & Tecnología:', 'IA generativa, Claude, API Integration, Agentes de IA, Machine Learning, People Analytics, casos de uso organizacionales (L&D, Product Strategy)'),
        ('Gestión de Proyectos:', 'Scrum Master (7+ años), Delivery Management, gestión de dependencias y riesgos, alineamiento de stakeholders'),
        ('Herramientas:', 'Jira, Azure DevOps, Confluence, Miro, herramientas de gestión de proyectos ágiles')
    ]
    
    for i, (categoria, items) in enumerate(competencias):
        is_last = (i == len(competencias) - 1)
        p = doc.add_paragraph()
        add_text_run(p, categoria, 9, bold=True)
        add_text_run(p, ' ' + items, CV_FONT_SIZES['body'])
        set_paragraph_format(p, space_after=5, is_last_in_section=is_last)
    
    # Línea separadora
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # ========== IDIOMAS (JUSTIFICADO) ==========
    
    header = doc.add_paragraph()
    add_text_run(header, 'IDIOMAS', CV_FONT_SIZES['header'], bold=True)
    set_paragraph_format(header, space_after=5)
    
    idiomas = doc.add_paragraph()
    add_text_run(idiomas, 'Español - Nativo | Inglés - Avanzado (C1)', CV_FONT_SIZES['body'])
    set_paragraph_format(idiomas, space_after=5, is_last_in_section=True)
    
    # Guardar documento usando ruta configurada
    output_path = get_output_path(output_filename)
    doc.save(str(output_path))

    return output_path

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='Generar CV en formato DOCX')
    parser.add_argument('-o', '--output',
                        default='CV_Gonzales_AttachGroup_AgileCoach.docx',
                        help='Nombre del archivo de salida')
    parser.add_argument('--nombre', help='Nombre del candidato')
    parser.add_argument('--ubicacion', help='Ubicación del candidato')
    parser.add_argument('--telefono', help='Teléfono del candidato')
    parser.add_argument('--email', help='Email del candidato')
    parser.add_argument('--linkedin', help='LinkedIn del candidato')

    args = parser.parse_args()

    # Preparar info del candidato si se proporcionaron argumentos
    candidate_info = None
    if any([args.nombre, args.ubicacion, args.telefono, args.email, args.linkedin]):
        candidate_info = CANDIDATE_INFO.copy()
        if args.nombre:
            candidate_info['nombre'] = args.nombre
        if args.ubicacion:
            candidate_info['ubicacion'] = args.ubicacion
        if args.telefono:
            candidate_info['telefono'] = args.telefono
        if args.email:
            candidate_info['email'] = args.email
        if args.linkedin:
            candidate_info['linkedin'] = args.linkedin

    output = create_cv_docx(output_filename=args.output, candidate_info=candidate_info)
    print(f"✅ CV generado exitosamente: {output}")
