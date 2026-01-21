#!/usr/bin/env python3
"""
Convertidor de CV Markdown a DOCX
Sistema: Easy Job Apply AI v2.0
Convierte cualquier CV en formato Markdown a DOCX con formato profesional ATS-friendly
"""

import sys
import re
import argparse
from pathlib import Path

# Agregar directorio scripts al path
sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from config import get_output_path, CV_MARGINS_CM, DEFAULT_FONT

def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=5, is_last_in_section=False):
    """Configura formato estándar del párrafo"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(10 if is_last_in_section else space_after)
    paragraph.alignment = alignment

def add_text_run(paragraph, text, font_size, bold=False, italic=False):
    """Helper para agregar texto con formato"""
    run = paragraph.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    return run

def parse_markdown_cv(md_file_path):
    """
    Parse del archivo Markdown para extraer secciones estructuradas

    Returns:
        dict con estructura del CV
    """
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    cv_data = {
        'nombre': '',
        'contacto': '',
        'secciones': []
    }

    current_section = None
    current_subsection = None

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Nombre (# TITULO)
        if line.startswith('# ') and not cv_data['nombre']:
            cv_data['nombre'] = line[2:].strip()
            continue

        # Contacto (línea con pipes y email)
        if '|' in line and '@' in line and not cv_data['contacto']:
            cv_data['contacto'] = line.replace('**', '').strip()
            continue

        # Separadores (---)
        if line.startswith('---'):
            continue

        # Headers de sección (##)
        if line.startswith('## '):
            section_name = line[3:].strip()
            current_section = {
                'tipo': 'seccion',
                'titulo': section_name,
                'contenido': []
            }
            cv_data['secciones'].append(current_section)
            current_subsection = None
            continue

        # Subsecciones (###)
        if line.startswith('### '):
            subsection_name = line[4:].strip()
            current_subsection = {
                'tipo': 'subseccion',
                'titulo': subsection_name,
                'contenido': []
            }
            if current_section:
                current_section['contenido'].append(current_subsection)
            continue

        # Texto en negrita que marca períodos o subtítulos (**texto**)
        if line.startswith('**') and line.endswith('**'):
            texto = line.replace('**', '').strip()
            item = {'tipo': 'periodo', 'texto': texto}
            if current_subsection:
                current_subsection['contenido'].append(item)
            elif current_section:
                current_section['contenido'].append(item)
            continue

        # Bullets (-)
        if line.startswith('- ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            item = {'tipo': 'bullet', 'texto': bullet_text}
            if current_subsection:
                current_subsection['contenido'].append(item)
            elif current_section:
                current_section['contenido'].append(item)
            continue

        # Checkmarks (✓)
        if line.startswith('✓ '):
            logro_text = line[2:].strip()
            item = {'tipo': 'logro', 'texto': logro_text}
            if current_section:
                current_section['contenido'].append(item)
            continue

        # Texto normal
        if line and current_section:
            item = {'tipo': 'texto', 'texto': line}
            if current_subsection:
                current_subsection['contenido'].append(item)
            else:
                current_section['contenido'].append(item)

    return cv_data

def convert_markdown_formatting(text):
    """
    Convierte formato Markdown a tuplas (texto, bold, italic)
    Retorna lista de segmentos: [(texto, bold, italic), ...]
    """
    segments = []
    current_pos = 0

    # Pattern para detectar **texto** (bold)
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')

    for match in bold_pattern.finditer(text):
        # Texto antes del bold
        if match.start() > current_pos:
            before_text = text[current_pos:match.start()]
            if before_text:
                segments.append((before_text, False, False))

        # Texto en bold
        segments.append((match.group(1), True, False))
        current_pos = match.end()

    # Texto restante
    if current_pos < len(text):
        remaining = text[current_pos:]
        if remaining:
            segments.append((remaining, False, False))

    return segments if segments else [(text, False, False)]

def add_formatted_paragraph(doc, text, font_size, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_after=5, is_last_in_section=False, default_bold=False):
    """
    Agrega párrafo con formato Markdown procesado
    """
    p = doc.add_paragraph()
    segments = convert_markdown_formatting(text)

    for segment_text, is_bold, is_italic in segments:
        add_text_run(p, segment_text, font_size, bold=is_bold or default_bold, italic=is_italic)

    set_paragraph_format(p, alignment=alignment, space_after=space_after,
                        is_last_in_section=is_last_in_section)
    return p

def create_docx_from_markdown(md_file_path, output_filename=None):
    """
    Convierte CV en Markdown a DOCX con formato profesional

    Args:
        md_file_path: Ruta al archivo Markdown
        output_filename: Nombre del archivo de salida (opcional)

    Returns:
        Path del archivo generado
    """
    # Parse del Markdown
    cv_data = parse_markdown_cv(md_file_path)

    # Crear documento
    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Cm(CV_MARGINS_CM)
    section.bottom_margin = Cm(CV_MARGINS_CM)
    section.left_margin = Cm(CV_MARGINS_CM)
    section.right_margin = Cm(CV_MARGINS_CM)

    # ========== HEADER: NOMBRE Y CONTACTO ==========
    if cv_data['nombre']:
        nombre = doc.add_paragraph()
        add_text_run(nombre, cv_data['nombre'], 18, bold=True)
        set_paragraph_format(nombre, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    if cv_data['contacto']:
        contacto = doc.add_paragraph()
        add_text_run(contacto, cv_data['contacto'], 9)
        set_paragraph_format(contacto, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5, is_last_in_section=True)

    # Separador inicial
    separador = doc.add_paragraph('_' * 90)
    set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # ========== SECCIONES ==========
    for i, seccion in enumerate(cv_data['secciones']):
        # Header de sección (## TITULO)
        header = doc.add_paragraph()
        add_text_run(header, seccion['titulo'], 12, bold=True)
        set_paragraph_format(header, space_after=5)

        # Procesar contenido de la sección
        process_section_content(doc, seccion['contenido'])

        # Separador después de cada sección (excepto la última)
        if i < len(cv_data['secciones']) - 1:
            separador = doc.add_paragraph('_' * 90)
            set_paragraph_format(separador, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)

    # Determinar nombre de salida
    if not output_filename:
        input_name = Path(md_file_path).stem
        output_filename = f"{input_name}.docx"

    # Guardar documento
    output_path = get_output_path(output_filename)
    doc.save(str(output_path))

    return output_path

def process_section_content(doc, contenido):
    """
    Procesa el contenido de una sección recursivamente
    """
    for idx, item in enumerate(contenido):
        is_last = (idx == len(contenido) - 1)

        if item['tipo'] == 'subseccion':
            # Subsección (### TITULO)
            subsec = doc.add_paragraph()
            add_text_run(subsec, item['titulo'], 9.5, bold=True)
            set_paragraph_format(subsec, space_after=5)

            # Procesar contenido de subsección
            process_section_content(doc, item['contenido'])

        elif item['tipo'] == 'periodo':
            # Período o subtítulo (**texto**)
            periodo = doc.add_paragraph()
            add_text_run(periodo, item['texto'], 9)
            set_paragraph_format(periodo, space_after=5)

        elif item['tipo'] == 'bullet':
            # Bullet point (- o •)
            add_formatted_paragraph(doc, '• ' + item['texto'], 9,
                                   space_after=5, is_last_in_section=is_last)

        elif item['tipo'] == 'logro':
            # Logro con checkmark (✓)
            add_formatted_paragraph(doc, '✓ ' + item['texto'], 9,
                                   space_after=5, is_last_in_section=is_last)

        elif item['tipo'] == 'texto':
            # Texto normal
            texto = item['texto']

            # Detectar si es encabezado en negrita seguido de contenido
            if ':' in texto and texto.strip().startswith('**'):
                # Formato: **Categoría:** contenido
                add_formatted_paragraph(doc, texto, 9, space_after=5)
            else:
                # Texto normal con posible formato Markdown
                add_formatted_paragraph(doc, texto, 9.5, space_after=5, is_last_in_section=is_last)

def main():
    parser = argparse.ArgumentParser(
        description='Convierte CV en formato Markdown a DOCX con formato profesional ATS-friendly'
    )
    parser.add_argument(
        'input',
        help='Archivo Markdown de entrada (ej: outputs/CV_Gonzales_Company_Position.md)'
    )
    parser.add_argument(
        '-o', '--output',
        help='Nombre del archivo DOCX de salida (opcional, por defecto usa el mismo nombre del input)'
    )

    args = parser.parse_args()

    # Validar que el archivo existe
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ Error: Archivo no encontrado: {args.input}")
        sys.exit(1)

    # Convertir
    try:
        output_path = create_docx_from_markdown(args.input, args.output)
        print(f"✅ CV generado exitosamente: {output_path}")
    except Exception as e:
        print(f"❌ Error al generar CV: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()
